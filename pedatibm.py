# Skrip Generator Rancangan Pengajaran PEDATI versi Bahasa Melayu dengan format kustom
import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.qn import qn
from io import BytesIO

# --- 1. KONFIGURASI ---
genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])

@st.cache_resource
def find_working_model():
    """Mencari model Gemini yang aktif secara otomatis."""
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                return m.name
    except:
        return "models/gemini-1.5-flash"
    return "models/gemini-1.5-flash"

selected_model_name = find_working_model()
model = genai.GenerativeModel(selected_model_name)


def generate_pedati_plan_bm(topic, syllabus, extra_context):
    # Prompt ini memaksa AI menggunakan Bahasa Melayu sepenuhnya dengan aturan ketat
    prompt = f"""
    Topik: {topic}. Kod Sukatan: {syllabus}. Konteks Tambahan: {extra_context}.
    Hasilkan satu rancangan pengajaran profesional dalam BAHASA MELAYU sepenuhnya.
    
    PERATURAN UTAMA PENGGUNAAN ISTILAH:
    1. JANGAN SEKALI-KALI menggunakan perkataan 'MURID'. Gantikan KESEMUANYA dengan istilah 'PELAJAR' di seluruh kandungan dokumen tanpa pengecualian.
    2. Gunakan nama fasa PEDATI yang tepat ini: P [Pengetahuan Sedia Ada], E [Empati / Penglibatan], D [Daya Usaha], A [Aplikasi], T [Taksiran], I [Impak / Penambahbaikan].

    PERATURAN KRITIKAL FORMAT KANDUNGAN:
    1. JANGAN gunakan simbol double asterisk (**) di mana-mana bahagian respons.
    2. JANGAN gunakan senarai bentuk bullet (seperti -, *, •). Anda WAJIB menggunakan senarai bernombor (1, 2, 3...) secara eksklusif untuk semua bahagian bersenarai.
    3. Semua penanda bahagian (SECTION:) di bawah mestilah ditulis dalam HURUF BESAR sepenuhnya.

    Strukturkan respons mengikut penanda bahagian (SECTION:) berikut:
    SECTION: TOPIK
    {topic}

    SECTION: OBJEKTIF PEMBELAJARAN
    [Sediakan tepat 4 isi menggunakan format bernombor 1., 2., 3., 4.]
    
    SECTION: HASIL PEMBELAJARAN
    [Sediakan tepat 4 isi menggunakan format bernombor 1., 2., 3., 4.]
    
    SECTION: KRITERIA KEJAYAAN
    [Sediakan tepat 4 isi menggunakan format bernombor 1., 2., 3., 4.]
    
    SECTION: PRASYARAT
    [Sediakan 1 pernyataan tentang pengetahuan sedia ada pelajar]
    
    SECTION: KATA KUNCI
    [Sediakan 6 item kata kunci penting yang dipisahkan oleh tanda koma sahaja. Jangan buat bentuk senarai.]
    
    SECTION: HOTS
    [Sediakan tepat 4 domain utama dalam Taksonomi Bloom menggunakan senarai bernombor]
    
    SECTION: KEWARGANEGARAAN DIGITAL
    [Sediakan tepat 4 isi bernombor 1., 2., 3., 4. tentang penggunaan sumber digital seperti YouTube, Canva, Chromebook, atau peranti digital oleh PELAJAR]

    SECTION: KANDUNGAN PEMBUKAAN PELAJARAN
    [Aktiviti set induksi/hook dan pelan transisi sebelum memulakan fasa PEDATI]

    SECTION: STRATEGI DIFERENSIASI (HIJAU)
    - HA (Higher Achiever): [1 aktiviti mencabar untuk pelajar berpencapaian tinggi]

    SECTION: STRATEGI DIFERENSIASI (KUNING)
    - MA (Medium Achiever): [1 aktiviti teras untuk pelajar berpencapaian sederhana]

    SECTION: STRATEGI DIFERENSIASI (MERAH)
    - LA (Lower Achiever): [1 aktiviti bersofisikasi/berpandu untuk pelajar berpencapaian rendah]

    SECTION: AKTIVITI PEMBELAJARAN TERADUN SATU (15 MINIT)
    - Aktiviti 1: [Penerangan aktiviti]
    - Persediaan Pensyarah: [Langkah demi langkah sebelum kelas bermula]
    - Objektif Aktiviti: [3 mata objektif]
    - Tugas Pelajar: [Butiran langkah demi langkah untuk PELAJAR]

    SECTION: AKTIVITI PEMBELAJARAN TERADUN DUA (15 MINIT)
    - Aktiviti 2: [Penerangan aktiviti]
    - Persediaan Pensyarah: [Langkah demi langkah sebelum kelas bermula]
    - Objektif Aktiviti: [3 mata objektif]
    - Tugas Pelajar: [Butiran langkah demi langkah untuk PELAJAR]

    SECTION: FASA PEDATI
    STAGE: P [Pengetahuan Sedia Ada] | SB: [Aktiviti Pensyarah] | CB: [Aktiviti Pelajar]
    STAGE: E [Empati / Penglibatan] | SB: [Aktiviti Pensyarah] | CB: [Aktiviti Pelajar]
    STAGE: D [Daya Usaha] | SB: [Aktiviti Pensyarah] | CB: [Aktiviti Pelajar]
    STAGE: A [Aplikasi] | SB: [Aktiviti Pensyarah] | CB: [Aktiviti Pelajar]
    STAGE: T [Taksiran] | SB: [Aktiviti Pensyarah] | CB: [Aktiviti Pelajar]
    STAGE: I [Impak / Penambahbaikan] | SB: [Aktiviti Pensyarah] | CB: [Aktiviti Pelajar]

    SECTION: PLENARI (TIKET KELUAR)
    [Aktiviti penutup ringkas sekitar 2-3 minit]

    SECTION: KERJA RUMAH
    [Tugasan susulan berdasarkan topik pelajaran]

    SECTION: CADANGAN TUGASAN MELANGKAH KE HADAPAN
    [Aktiviti hook dan pelan transisi untuk sesi pembelajaran hari esok]
    """
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"System Error: {str(e)}"


def add_page_number(run):
    """Fungsi pembantu untuk memasukkan nombor halaman dinamik Word di bahagian atas tengah."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def create_word_export(topic, syllabus, text):
    doc = Document()
    
    # Peraturan 1: Tetapkan Saiz Kertas LETTER dan Margin 0.5 inci di keempat-empat penjuru
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        # Peraturan 7: Letakkan Nombor Halaman di Header, Bahagian Atas Tengah
        header = section.header
        header_p = header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_p.add_run()
        header_run.font.name = 'Arial'
        header_run.font.size = Pt(10)
        add_page_number(header_run)

    # Peraturan 2 & 8: Tajuk Utama dalam HURUF BESAR dengan Saiz Font 14
    main_title = f'RANCANGAN PENGAJARAN: {topic} ({syllabus})'.upper()
    title_p = doc.add_heading(level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(main_title)
    title_run.font.size = Pt(14)
    title_run.bold = True
    
    # Peraturan 3 & 8: Tetapkan Font lalai Arial Saiz 12 dan Jarak Perenggan TUNGGAL (Single)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # 1. Jadual Pentadbiran (Admin Table)
    admin_table = doc.add_table(rows=3, cols=4)
    admin_table.style = 'Table Grid'
    labels = [["Minggu Ke:", "Tarikh:"], ["Bilangan Pelajar:", "Hari:"], ["Tempat / No Makmal:", "Durasi (minit):"]]
    for r in range(3):
        admin_table.cell(r, 0).text = labels[r][0]
        admin_table.cell(r, 2).text = labels[r][1]
    
    for row in admin_table.rows:
        for cell in row.cells:
            for d_p in cell.paragraphs:
                d_p.paragraph_format.line_spacing = 1.0
                for run in d_p.runs:
                    run.font.size = Pt(12)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.line_spacing = 1.0

    # 2. Jadual Sumber & Bahan
    r_heading = doc.add_paragraph()
    r_heading.paragraph_format.line_spacing = 1.0
    r_run = r_heading.add_run("SUMBER & BAHAN PELAJARAN")
    r_run.bold = True
    r_run.font.size = Pt(14)
    
    res_table = doc.add_table(rows=1, cols=1)
    res_table.style = 'Table Grid'
    res_table.cell(0, 0).text = "Papan pintar (Smart board), Chromebook, Meja tulis, Projektor, Perkongsian skrin komputer riba"
    res_table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = 1.0

    # 3. Proses Pembahagian Kandungan (Parsing) & Kotak Jadual
    sections = text.split('SECTION:')
    for section in sections:
        if not section.strip(): continue
        lines = section.strip().split('\n')
        
        # Peraturan 2 & 4: Paksa Huruf Besar untuk Tajuk Bahagian & buang asterisk
        title = lines[0].strip().replace("**", "").upper()
        content_lines = lines[1:]
        
        doc_heading = doc.add_paragraph()
        doc_heading.paragraph_format.line_spacing = 1.0
        h_run = doc_heading.add_run(title)
        h_run.bold = True
        h_run.font.size = Pt(14)  # Peraturan 8: Saiz font tajuk 14

        # Peraturan 6: Pengurusan grid KATA KUNCI kustom
        if "KATA KUNCI" in title or "KEYWORDS" in title:
            raw_keywords_text = " ".join([l.strip() for l in content_lines if l.strip()])
            keyword_items = [kw.strip() for kw in raw_keywords_text.split(",") if kw.strip()]
            
            # Membina jadual grid 2 baris x 3 lajur
            kw_table = doc.add_table(rows=2, cols=3)
            kw_table.style = 'Table Grid'
            
            idx = 0
            for r in range(2):
                for c in range(3):
                    if idx < len(keyword_items):
                        cell = kw_table.cell(r, c)
                        cell.text = keyword_items[idx]
                        # Diselaraskan ke tengah (Aligned center)
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.line_spacing = 1.0
                        if p.runs:
                            p.runs[0].font.size = Pt(12)
                        idx += 1
            doc.add_paragraph().paragraph_format.line_spacing = 1.0
            
        elif "|" in section and "FASA" in title:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = 'Fasa (PEDATI)', 'Aktiviti Pensyarah', 'Aktiviti Pelajar'
            
            for cell in hdr:
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing = 1.0
                if p.runs:
                    p.runs[0].font.size = Pt(12)
                    p.runs[0].font.bold = True

            for line in content_lines:
                if "|" in line:
                    p_split = line.split("|")
                    row = table.add_row().cells
                    row[0].text = p_split[0].split(":")[-1].strip().replace("**", "")
                    row[1].text = p_split[1].split(":")[-1].strip().replace("**", "")
                    row[2].text = p_split[2].split(":")[-1].strip().replace("**", "")
                    
                    for cell in row:
                        p = cell.paragraphs[0]
                        p.paragraph_format.line_spacing = 1.0
                        if p.runs:
                            p.runs[0].font.size = Pt(12)
            doc.add_paragraph().paragraph_format.line_spacing = 1.0
        else:
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            
            cleaned_content = "\n".join([l.strip() for l in content_lines if l.strip()]).replace("**", "")
            table.cell(0, 0).text = cleaned_content
            
            p = table.cell(0, 0).paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            if p.runs:
                p.runs[0].font.size = Pt(12)
            doc.add_paragraph().paragraph_format.line_spacing = 1.0

    # 4. Halaman Pengesahan HOD
    doc.add_page_break()
    
    hod_heading = doc.add_paragraph()
    hod_heading.paragraph_format.line_spacing = 1.0
    hod_run = hod_heading.add_run("ULASAN & PENGESAHAN HOD")
    hod_run.bold = True
    hod_run.font.size = Pt(14)
    
    hod_table = doc.add_table(rows=3, cols=2)
    hod_table.style = 'Table Grid'
    hod_table.cell(0, 0).text = "Ulasan / Catatan"
    hod_table.cell(0, 1).text = "Tandatangan / Cop Rasmi"
    hod_table.rows[1].height = Pt(60)
    hod_table.cell(2, 0).text = "Tarikh:"
    hod_table.cell(2, 1).text = "Nama:"
    
    for row in hod_table.rows:
        for cell in row.cells:
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            if p.runs:
                p.runs[0].font.size = Pt(12)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# --- 4. BAHAGIAN GUI STREAMLIT ---
st.set_page_config(page_title="Penjana Master PEDATI", layout="wide")

with st.sidebar:
    st.title("📖 Panduan Pengguna")
    st.info("Cara menggunakan portal ini:")
    st.markdown("""
    ### 1. Isi Maklumat
    Masukkan **Topik Pelajaran** dan **Kod Sukatan**. Gunakan kotak **Konteks** untuk keperluan khusus seperti "Kerja Berkumpulan", "Kajian Kes", atau pautan YouTube.
    
    ### 2. Jana Pelan
    Klik butang **🚀 JANA** dan tunggu AI merangka rancangan pengajaran anda.
    
    ### 3. Muat Turun
    Semak draf pada **Pratonton AI**. Jika berpuas hati, klik **📥 Muat Turun Word** untuk mendapatkan fail rasmi anda.
    """)
    st.markdown("---")
    st.caption("Versi Aplikasi 4.0 | Enjin Grid Kustom")

st.title("🎓 Penjana Rancangan Pengajaran PEDATI (Versi BM)")
st.info(f"Sistem dihubungkan melalui core: {selected_model_name}")

c1, c2 = st.columns(2)
with c1: u_topic = st.text_input("Topik Pelajaran:")
with c2: u_syllabus = st.text_input("Kod Sukatan Pelajaran:")
u_extra = st.text_area("Konteks/Kata Kunci Spesifik (Pilihan):")

if st.button("🚀 JANA RANCANGAN PENGAJARAN PEDATI"):
    if u_topic and u_syllabus:
        with st.spinner("AI sedang membina rancangan pengajaran struktur PEDATI..."):
            result = generate_pedati_plan_bm(u_topic, u_syllabus, u_extra)
            # Membersihkan sebarang asterisks sisa dari paparan web
            st.session_state['pedati_out_bm'] = result.replace("**", "")
    else:
        st.warning("Sila isi ruang Topik dan Kod Sukatan Pelajaran.")

if 'pedati_out_bm' in st.session_state:
    st.divider()
    st.subheader("📝 Pratonton Draf AI (Format Bersih)")
    st.text_area("Kandungan", st.session_state['pedati_out_bm'], height=350)
    
    doc_file = create_word_export(u_topic, u_syllabus, st.session_state['pedati_out_bm'])
    st.download_button("📥 Muat Turun ke Versi Word (.docx)", doc_file, f"Rancangan_PEDATI_{u_topic}.docx")

st.markdown("---") 
st.markdown(
    """
    <div style='text-align: center; color: grey; font-size: 0.8em;'>
        <p><b>Aplikasi Kecerdasan Buatan Rancangan Pengajaran PEDATI v4.0</b></p>
        <p>Dibangunkan & Dikonsepkan oleh: <b>Hajah Nurul Haziqah @ Hjh Hartini Hj Nordin</b></p>
        <p>© 2026 PTES Academic Innovation Computer Science</p>
    </div>
    """,
    unsafe_allow_html=True
)
